#!/usr/bin/env python3
"""
Professional Resume Generator
Converts portfolio website content into beautifully formatted PDF and DOCX resumes
"""

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

class ResumeGenerator:
    def __init__(self):
        self.data = self.get_resume_data()
        
    def get_resume_data(self):
        """Extract resume data from portfolio content"""
        return {
            'name': 'Michael Fedorovsky',
            'title': 'DevOps Engineer',
            'contact': {
                'email': 'michaelovsky5@gmail.com',
                'phone': '054-763-2418',
                'location': 'Bat Yam, Israel',
                'github': 'https://github.com/Michaelunkai',
                'linkedin': 'https://linkedin.com/in/michael-fedorovsky-b26099278'
            },
            'summary': """DevOps Engineer with 3+ years of hands-on experience building and maintaining production infrastructure that powers real businesses. Expert in cloud infrastructure, CI/CD pipelines, containerization, and monitoring. Built 50+ production-grade automation tools achieving 99.9% uptime across multi-cloud environments. Specialized in cloud automation, CI/CD pipelines, container orchestration, and cost optimization. Delivered 50+ production tools, reduced cloud costs by 40%, and maintained 99.9% uptime across AWS, Azure, and GCP.""",
            
            'experience': [
                {
                    'company': 'TovTech',
                    'position': 'DevOps Engineer',
                    'period': '2021 - Present',
                    'achievements': [
                        'Designed and deployed multi-environment cloud infrastructure on AWS and Webdock with Cloudflare CDN integration',
                        'Achieved 99.9% uptime and 60% faster deployment times',
                        'Built comprehensive CI/CD pipelines with GitHub Actions, automated testing workflows, Docker multi-stage builds',
                        'Implemented monitoring stack (Prometheus + Grafana + ELK) for real-time system health and security visibility',
                        'Reduced cloud costs by 30-60% through right-sizing, reserved/spot instances, and idle resource detection',
                        'Developed 50+ production-grade automation tools including container security scanners and API gateways'
                    ]
                }
            ],
            
            'projects': [
                {
                    'name': 'TovPlay Gaming Platform',
                    'description': 'Full-Stack Gaming Backend with Real-Time Features',
                    'details': 'Production gaming backend built with Flask, PostgreSQL, Socket.IO, and Docker. Features JWT authentication, Discord OAuth integration, real-time multiplayer capabilities, and complete CI/CD pipeline. Handles 1000+ concurrent users.',
                    'tech': ['Flask', 'PostgreSQL', 'Socket.IO', 'Docker', 'GitHub Actions', 'JWT'],
                    'url': 'https://github.com/Michaelunkai/TovPlay'
                },
                {
                    'name': 'Game Library Manager',
                    'description': 'Docker-Based Game Library with Web UI - 928 Games',
                    'details': 'Dockerized game library management system with modern web interface. Browse, search, and manage 928 games with detailed metadata, cover art, and download links.',
                    'tech': ['Docker', 'JavaScript', 'Node.js', 'Web UI'],
                    'url': 'https://github.com/Michaelunkai/game-library-manager-web'
                },
                {
                    'name': 'Container Security Scanner',
                    'description': 'Automated CVE Detection & CVSS Scoring',
                    'details': 'Automated vulnerability scanning tool for Docker images using Trivy. Provides severity reporting, integrates with CI/CD pipelines, and generates compliance reports.',
                    'tech': ['Trivy', 'Docker', 'Python', 'GitHub Actions', 'Shell'],
                    'url': 'https://github.com/Michaelunkai/container-security-scanner'
                }
            ],
            
            'skills': {
                'Cloud & Infrastructure': ['AWS', 'Azure', 'Google Cloud', 'DigitalOcean', 'Cloudflare', 'Vercel', 'Netlify'],
                'Container & Orchestration': ['Docker', 'Kubernetes', 'Helm', 'ArgoCD', 'Rancher', 'Podman', 'containerd'],
                'Infrastructure as Code': ['Terraform', 'Ansible', 'Pulumi', 'CloudFormation', 'Vagrant', 'Chef', 'Puppet'],
                'CI/CD': ['GitHub Actions', 'GitLab CI', 'Jenkins', 'CircleCI', 'Travis CI', 'TeamCity', 'Drone'],
                'Monitoring & Observability': ['Prometheus', 'Grafana', 'Datadog', 'New Relic', 'Jaeger', 'Zipkin', 'Sentry'],
                'Logging & Analytics': ['Elasticsearch', 'Logstash', 'Kibana', 'Fluentd', 'Loki', 'Splunk'],
                'Databases': ['PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Cassandra', 'DynamoDB', 'MariaDB'],
                'Message Queues': ['RabbitMQ', 'Apache Kafka', 'NATS', 'ActiveMQ', 'Amazon SQS'],
                'Web Servers & Load Balancers': ['Nginx', 'Apache', 'HAProxy', 'Traefik', 'Caddy', 'Envoy'],
                'Programming & Scripting': ['Python', 'Bash', 'PowerShell', 'Go', 'Node.js', 'TypeScript', 'C#'],
                'Version Control': ['Git', 'GitHub', 'GitLab', 'Bitbucket'],
                'Operating Systems': ['Linux', 'Ubuntu', 'CentOS', 'Debian', 'Red Hat', 'Alpine Linux', 'Windows Server'],
                'Security': ['HashiCorp Vault', 'Trivy', 'Snyk', 'SonarQube', 'OWASP ZAP', 'Aqua Security'],
                'Service Mesh': ['Istio', 'Linkerd', 'Consul'],
                'APIs': ['REST APIs', 'GraphQL', 'gRPC', 'Swagger', 'Postman'],
                'Testing': ['Jest', 'Pytest', 'Selenium', 'Cypress', 'JUnit'],
                'Build Tools': ['Maven', 'Gradle', 'npm', 'Yarn', 'Make', 'Helm'],
                'Package Managers': ['apt', 'yum', 'Chocolatey', 'Homebrew'],
                'Serverless': ['AWS Lambda', 'Azure Functions', 'Google Cloud Functions', 'Serverless Framework'],
                'Frameworks': ['Flask', 'FastAPI', 'Express.js', '.NET', 'Spring Boot'],
                'Virtualization': ['VMware', 'VirtualBox', 'KVM', 'Hyper-V']
            },
            
            'certifications': [
                'OWASP Security Best Practices',
                'GDPR Compliance',
                'SOC2 Compliance'
            ],
            
            'achievements': [
                '50+ production-grade automation tools built',
                '99.9% uptime achieved across multi-cloud environments',
                '3+ years of DevOps experience',
                '50+ GitHub repositories',
                '30-60% cloud cost reduction',
                '1000+ concurrent users handled'
            ]
        }
    
    def generate_pdf(self, filename='F:/study/resume/Michael_Fedorovsky_Resume.pdf'):
        """Generate beautifully formatted PDF resume"""
        doc = SimpleDocTemplate(filename, pagesize=letter,
                                rightMargin=0.75*inch, leftMargin=0.75*inch,
                                topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define custom styles
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=6,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Subtitle style
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#3b82f6'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica'
        )
        
        # Section heading style
        section_style = ParagraphStyle(
            'SectionHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1e40af'),
            spaceAfter=8,
            spaceBefore=12,
            fontName='Helvetica-Bold',
            borderWidth=1,
            borderColor=colors.HexColor('#3b82f6'),
            borderPadding=4,
            backColor=colors.HexColor('#eff6ff')
        )
        
        # Contact info style
        contact_style = ParagraphStyle(
            'Contact',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#4b5563'),
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        # Body text style
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#374151'),
            spaceAfter=6,
            alignment=TA_JUSTIFY
        )
        
        # Bullet style
        bullet_style = ParagraphStyle(
            'Bullet',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#374151'),
            leftIndent=20,
            spaceAfter=4
        )
        
        # Add content
        # Header
        elements.append(Paragraph(self.data['name'], title_style))
        elements.append(Paragraph(self.data['title'], subtitle_style))
        
        # Contact information
        contact_text = f"""
        {self.data['contact']['email']} | {self.data['contact']['phone']} | {self.data['contact']['location']}<br/>
        GitHub: {self.data['contact']['github']} | LinkedIn: {self.data['contact']['linkedin']}
        """
        elements.append(Paragraph(contact_text, contact_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Professional Summary
        elements.append(Paragraph('PROFESSIONAL SUMMARY', section_style))
        elements.append(Paragraph(self.data['summary'], body_style))
        elements.append(Spacer(1, 0.15*inch))
        
        # Key Achievements
        elements.append(Paragraph('KEY ACHIEVEMENTS', section_style))
        for achievement in self.data['achievements']:
            elements.append(Paragraph(f'• {achievement}', bullet_style))
        elements.append(Spacer(1, 0.15*inch))
        
        # Professional Experience
        elements.append(Paragraph('PROFESSIONAL EXPERIENCE', section_style))
        for exp in self.data['experience']:
            # Company and position
            exp_header = f"<b>{exp['position']}</b> - {exp['company']} | {exp['period']}"
            elements.append(Paragraph(exp_header, body_style))
            elements.append(Spacer(1, 0.05*inch))
            
            # Achievements
            for achievement in exp['achievements']:
                elements.append(Paragraph(f'• {achievement}', bullet_style))
            elements.append(Spacer(1, 0.15*inch))
        
        # Projects
        elements.append(Paragraph('FEATURED PROJECTS', section_style))
        for project in self.data['projects']:
            # Project name and description
            proj_header = f"<b>{project['name']}</b> - {project['description']}"
            elements.append(Paragraph(proj_header, body_style))
            elements.append(Paragraph(project['details'], body_style))
            
            # Technologies
            tech_text = '<b>Technologies:</b> ' + ', '.join(project['tech'])
            elements.append(Paragraph(tech_text, body_style))
            elements.append(Paragraph(f"<b>URL:</b> {project['url']}", body_style))
            elements.append(Spacer(1, 0.1*inch))
        
        # Technical Skills
        elements.append(PageBreak())
        elements.append(Paragraph('TECHNICAL SKILLS', section_style))
        
        for category, skills in self.data['skills'].items():
            skill_text = f"<b>{category}:</b> {', '.join(skills[:10])}"  # Limit to 10 per line
            if len(skills) > 10:
                skill_text += f"<br/>{', '.join(skills[10:])}"
            elements.append(Paragraph(skill_text, body_style))
            elements.append(Spacer(1, 0.05*inch))
        
        # Build PDF
        doc.build(elements)
        return filename
    
    def generate_docx(self, filename='F:/study/resume/Michael_Fedorovsky_Resume.docx'):
        """Generate beautifully formatted DOCX resume"""
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading(self.data['name'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(30, 64, 175)
        title_run.font.size = Pt(24)
        
        # Subtitle
        subtitle = doc.add_heading(self.data['title'], level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = RGBColor(59, 130, 246)
        subtitle_run.font.size = Pt(16)
        
        # Contact information
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact.add_run(
            f"{self.data['contact']['email']} | {self.data['contact']['phone']} | {self.data['contact']['location']}\n"
            f"GitHub: {self.data['contact']['github']} | LinkedIn: {self.data['contact']['linkedin']}"
        )
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = RGBColor(75, 85, 99)
        
        doc.add_paragraph()  # Spacer
        
        # Professional Summary
        self.add_section_heading(doc, 'PROFESSIONAL SUMMARY')
        summary_para = doc.add_paragraph(self.data['summary'])
        summary_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Key Achievements
        self.add_section_heading(doc, 'KEY ACHIEVEMENTS')
        for achievement in self.data['achievements']:
            p = doc.add_paragraph(achievement, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        
        # Professional Experience
        self.add_section_heading(doc, 'PROFESSIONAL EXPERIENCE')
        for exp in self.data['experience']:
            # Company and position
            exp_para = doc.add_paragraph()
            exp_run = exp_para.add_run(f"{exp['position']} - {exp['company']}")
            exp_run.bold = True
            exp_run.font.size = Pt(12)
            
            period_run = exp_para.add_run(f" | {exp['period']}")
            period_run.italic = True
            period_run.font.size = Pt(11)
            
            # Achievements
            for achievement in exp['achievements']:
                p = doc.add_paragraph(achievement, style='List Bullet')
                p.paragraph_format.left_indent = Inches(0.5)
        
        # Projects
        self.add_section_heading(doc, 'FEATURED PROJECTS')
        for project in self.data['projects']:
            # Project name
            proj_para = doc.add_paragraph()
            proj_name = proj_para.add_run(project['name'])
            proj_name.bold = True
            proj_name.font.size = Pt(12)
            
            proj_desc = proj_para.add_run(f" - {project['description']}")
            proj_desc.font.size = Pt(11)
            
            # Details
            doc.add_paragraph(project['details'])
            
            # Technologies
            tech_para = doc.add_paragraph()
            tech_label = tech_para.add_run('Technologies: ')
            tech_label.bold = True
            tech_para.add_run(', '.join(project['tech']))
            
            # URL
            url_para = doc.add_paragraph()
            url_label = url_para.add_run('URL: ')
            url_label.bold = True
            url_para.add_run(project['url'])
            
            doc.add_paragraph()  # Spacer
        
        # Technical Skills
        doc.add_page_break()
        self.add_section_heading(doc, 'TECHNICAL SKILLS')
        
        for category, skills in self.data['skills'].items():
            skill_para = doc.add_paragraph()
            cat_run = skill_para.add_run(f'{category}: ')
            cat_run.bold = True
            cat_run.font.size = Pt(11)
            
            skill_run = skill_para.add_run(', '.join(skills))
            skill_run.font.size = Pt(10)
        
        # Save document
        doc.save(filename)
        return filename
    
    def add_section_heading(self, doc, text):
        """Add a styled section heading to Word document"""
        heading = doc.add_heading(text, level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(30, 64, 175)
        heading_run.font.size = Pt(14)
        heading_run.bold = True
        
        # Add a horizontal line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)

# Main execution
if __name__ == '__main__':
    print("Generating professional resume files...")
    
    generator = ResumeGenerator()
    
    # Generate PDF
    print("Creating PDF resume...")
    pdf_file = generator.generate_pdf()
    print(f"PDF created: {pdf_file}")
    
    # Generate DOCX
    print("Creating DOCX resume...")
    docx_file = generator.generate_docx()
    print(f"DOCX created: {docx_file}")
    
    print("\nResume generation complete!")
    print(f"\nFiles saved in: F:/study/resume/")
    print(f"   - PDF: Michael_Fedorovsky_Resume.pdf")
    print(f"   - DOCX: Michael_Fedorovsky_Resume.docx")
